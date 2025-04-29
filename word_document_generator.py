from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class WordDocumentGenerator:
    def __init__(self, gedcom_processor):
        self.gedcom_processor = gedcom_processor

    def set_cell_vertical_text(self, cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        textDirection = OxmlElement('w:textDirection')
        textDirection.set(qn('w:val'), 'btLr')
        tcPr.append(textDirection)

    def generate_word_grid(self, doc: Document, root_person, depth: int = 4):
        # Set landscape orientation
        section = doc.sections[0]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(29.7)  # A4 height
        section.page_height = Cm(21.0)  # A4 width

        table = doc.add_table(rows=depth, cols=2 ** (depth - 1))
        table.style = 'Table Grid'
        table.allow_autofit = True
        table.width = Inches(11)

        for generation in range(depth):
            cells = [""] * (2 ** generation)
            stack = [(root_person, 0)]
            cell_index = 0

            while stack and cell_index < len(cells):
                person, gen = stack.pop(0)
                if gen == generation:
                    if person:
                        cells[cell_index] = self.gedcom_processor.get_person_info(person)
                    cell_index += 1
                elif person:
                    father, mother = self.gedcom_processor.find_parents(person)
                    stack.append((mother, gen + 1))
                    stack.append((father, gen + 1))

            row = table.rows[generation]
            cells_per_column = (2 ** (depth - 1)) // (2 ** generation)
            for i, cell_text in enumerate(cells):
                start_col = i * cells_per_column
                end_col = (i + 1) * cells_per_column
                if start_col < end_col:
                    if cells_per_column > 1:
                        cell = row.cells[start_col]
                        for col in range(start_col + 1, end_col):
                            cell.merge(row.cells[col])

                    cell = row.cells[start_col]
                    cell.text = cell_text

                    if generation == depth - 1:
                        self.set_cell_vertical_text(cell)
                        tc = cell._tc
                        tc.height = Cm(8)

                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
                    font_size = 12 - generation
                    run.font.size = Pt(max(8, font_size))

        doc.add_paragraph('')